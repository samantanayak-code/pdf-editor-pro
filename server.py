"""
PDF Backend Server — http://localhost:5050
Routes:
  GET  /api/health
  POST /api/preview/word    → {html, scanned}   fast, no OCR
  POST /api/convert/word    → .docx download     full quality
  POST /api/preview/excel   → {tables:[...]}     fast
  POST /api/convert/excel   → .xlsx download
"""

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import pdfplumber, openpyxl, io, os, re, html as html_module
from collections import defaultdict
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

app = Flask(__name__)
CORS(app)

# ─────────────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _is_bold(fontname: str) -> bool:
    fn = (fontname or "").lower()
    return "bold" in fn or "heavy" in fn or "black" in fn

def _is_italic(fontname: str) -> bool:
    fn = (fontname or "").lower()
    return "italic" in fn or "oblique" in fn

def _clean(v):
    return re.sub(r'\s+', ' ', str(v)).strip() if v else ""

def _is_scanned(pdf_bytes: bytes) -> bool:
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            return sum(len(p.extract_text() or "") for p in pdf.pages) < 30
    except:
        return True

def _setup_tesseract():
    import platform, shutil, pytesseract
    if platform.system() == "Windows" and not shutil.which("tesseract"):
        for c in [r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                  r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"]:
            if os.path.isfile(c):
                pytesseract.pytesseract.tesseract_cmd = c; break


# ─────────────────────────────────────────────────────────────────────────────
# CORE: Extract page content with rich formatting
# Returns list of "blocks": either {type:"para", runs:[...]} or {type:"table", rows:[[cells]]}
# Each run: {text, bold, italic, size, align}
# Each cell: list of runs
# ─────────────────────────────────────────────────────────────────────────────

def _char_props(char):
    return {
        "bold": _is_bold(char.get("fontname", "")),
        "italic": _is_italic(char.get("fontname", "")),
        "size": round(char.get("size", 10)),
    }

def _words_with_style(page, bbox=None):
    """Return words within bbox (or whole page) with bold/italic/size from chars."""
    crop = page.within_bbox(bbox) if bbox else page
    words = crop.extract_words(extra_attrs=["fontname", "size"])
    result = []
    for w in words:
        # Determine dominant font from chars in this word's bbox
        try:
            wchars = [c for c in crop.chars
                      if c["x0"] >= w["x0"]-1 and c["x1"] <= w["x1"]+1
                      and c["top"] >= w["top"]-1 and c["bottom"] <= w["bottom"]+1
                      and c["text"].strip()]
        except:
            wchars = []
        if wchars:
            bold = any(_is_bold(c.get("fontname","")) for c in wchars)
            italic = any(_is_italic(c.get("fontname","")) for c in wchars)
            size = round(sum(c.get("size",10) for c in wchars) / len(wchars))
        else:
            bold = _is_bold(w.get("fontname",""))
            italic = _is_italic(w.get("fontname",""))
            size = round(w.get("size", 10))
        result.append({
            "text": w["text"],
            "x0": w["x0"], "x1": w["x1"],
            "top": w["top"], "bottom": w["bottom"],
            "bold": bold, "italic": italic, "size": size,
        })
    return result

def _group_words_into_lines(words, y_tol=3):
    """Group words into lines by similar top-y, then sort left-to-right."""
    if not words:
        return []
    lines = defaultdict(list)
    for w in words:
        key = round(w["top"] / y_tol) * y_tol
        lines[key].append(w)
    result = []
    for key in sorted(lines.keys()):
        line = sorted(lines[key], key=lambda w: w["x0"])
        result.append(line)
    return result

def _lines_to_runs(lines, page_width=None):
    """Convert lines-of-words into paragraph run objects with alignment."""
    paragraphs = []
    for line in lines:
        if not line:
            continue
        # Detect alignment from x position
        left_x = line[0]["x0"]
        right_x = line[-1]["x1"]
        center_x = (left_x + right_x) / 2
        align = WD_ALIGN_PARAGRAPH.LEFT
        if page_width:
            mid = page_width / 2
            if abs(center_x - mid) < 40 and left_x > page_width * 0.2:
                align = WD_ALIGN_PARAGRAPH.CENTER

        # Build runs, merging consecutive words with same style
        runs = []
        for w in line:
            if runs and runs[-1]["bold"] == w["bold"] and runs[-1]["italic"] == w["italic"] and runs[-1]["size"] == w["size"]:
                runs[-1]["text"] += " " + w["text"]
            else:
                if runs:
                    runs[-1]["text"] += " "  # trailing space before new run
                runs.append({"text": w["text"], "bold": w["bold"], "italic": w["italic"], "size": w["size"]})
        paragraphs.append({"runs": runs, "align": align})
    return paragraphs

def _extract_cell_content(page, bbox):
    """Extract rich-formatted content from a table cell bbox."""
    words = _words_with_style(page, bbox)
    lines = _group_words_into_lines(words)
    return _lines_to_runs(lines)  # list of {runs, align}

def _extract_page_blocks(page):
    """
    Extract all content on a page as blocks.
    Returns list of:
      {"type": "para",  "align": ..., "runs": [{text, bold, italic, size}]}
      {"type": "table", "col_widths": [...], "rows": [[cell_paras, ...], ...]}
    """
    blocks = []
    pw = page.width

    # Find tables
    found_tables = page.find_tables()
    table_bboxes = [ft.bbox for ft in found_tables]

    def in_any_table(word):
        for bb in table_bboxes:
            if (word["x0"] >= bb[0]-2 and word["x1"] <= bb[2]+2
                    and word["top"] >= bb[1]-2 and word["bottom"] <= bb[3]+2):
                return True
        return False

    # Words outside tables
    all_words = _words_with_style(page)
    outside_words = [w for w in all_words if not in_any_table(w)]
    para_lines = _group_words_into_lines(outside_words)
    # Sort para_lines and tables together by Y position
    items = []  # (y_pos, block)
    para_groups = _lines_to_runs(para_lines, pw)
    # Track y positions of lines
    line_y = {}
    for i, line in enumerate(para_lines):
        if line:
            line_y[i] = line[0]["top"]

    for i, pg in enumerate(para_groups):
        y = line_y.get(i, 0)
        items.append((y, {"type": "para", **pg}))

    for ft in found_tables:
        raw_table = ft.extract()
        if not raw_table:
            continue
        # Get column widths from table cells bounding boxes
        # ft.bbox = (x0, top, x1, bottom)
        tb_x0, tb_top, tb_x1, tb_bottom = ft.bbox
        # Derive column x-boundaries from the raw cells
        try:
            col_xs = [cell.bbox[0] for cell in ft.cells[0]] + [ft.cells[0][-1].bbox[2]]
            col_widths_pts = [col_xs[i+1]-col_xs[i] for i in range(len(col_xs)-1)]
        except:
            n_cols = max(len(r) for r in raw_table)
            col_width = (tb_x1 - tb_x0) / max(n_cols, 1)
            col_widths_pts = [col_width] * n_cols

        # Extract each cell with formatting
        rows_out = []
        try:
            for row_cells in ft.cells:
                row_out = []
                for cell in row_cells:
                    cell_bbox = cell.bbox  # (x0, top, x1, bottom)
                    paras = _extract_cell_content(page, cell_bbox)
                    row_out.append(paras)
                rows_out.append(row_out)
        except Exception:
            # Fallback: use raw text extraction
            for raw_row in raw_table:
                row_out = []
                for cell_txt in raw_row:
                    txt = _clean(cell_txt)
                    row_out.append([{"runs": [{"text": txt, "bold": False, "italic": False, "size": 9}], "align": WD_ALIGN_PARAGRAPH.LEFT}] if txt else [])
                rows_out.append(row_out)

        items.append((tb_top, {"type": "table", "col_widths_pts": col_widths_pts, "rows": rows_out}))

    items.sort(key=lambda x: x[0])
    return [b for _, b in items]


# ─────────────────────────────────────────────────────────────────────────────
# WORD DOCUMENT BUILDER
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_borders(cell, color="AAAAAA"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for s in ("top","left","bottom","right"):
        b = OxmlElement(f"w:{s}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), color)
        tcB.append(b)
    tcPr.append(tcB)

def _set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _add_runs_to_para(para, runs, default_size=10):
    for run_data in runs:
        run = para.add_run(run_data["text"])
        run.bold = run_data.get("bold", False)
        run.italic = run_data.get("italic", False)
        sz = run_data.get("size", default_size)
        run.font.size = Pt(sz if sz > 0 else default_size)
        run.font.name = "Arial"

def _add_para_block(doc, block, default_size=10, space_before=0, space_after=3):
    runs = block.get("runs", [])
    if not any(r.get("text","").strip() for r in runs):
        return
    p = doc.add_paragraph()
    p.alignment = block.get("align", WD_ALIGN_PARAGRAPH.LEFT)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(13)
    _add_runs_to_para(p, runs, default_size)
    return p

def _pts_to_dxa(pts):
    """Convert PDF points to Word DXA (1/20 of a point → twips). 1 pt = 20 twips."""
    return int(pts * 20)

def _add_table_block(doc, block, page_width_pts=595.32, margin_pts=56):
    rows = block["rows"]
    col_widths_pts = block["col_widths_pts"]
    if not rows or not col_widths_pts:
        return

    n_cols = len(col_widths_pts)
    n_rows = len(rows)

    # Scale column widths to fit page content width
    content_width_pts = page_width_pts - 2 * margin_pts
    total_pts = sum(col_widths_pts)
    if total_pts > 0:
        scale = content_width_pts / total_pts
        col_widths_dxa = [int(w * scale * 20) for w in col_widths_pts]
    else:
        w = int(content_width_pts * 20 / n_cols)
        col_widths_dxa = [w] * n_cols

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"
    # Set column widths
    for ci, w in enumerate(col_widths_dxa):
        for row in table.rows:
            row.cells[ci].width = w

    for ri, row_data in enumerate(rows):
        is_header = (ri == 0)
        for ci in range(n_cols):
            cell = table.cell(ri, ci)
            cell.paragraphs[0].clear()
            # Remove default paragraph
            for existing_para in cell.paragraphs[1:]:
                p = existing_para._element
                p.getparent().remove(p)

            cell_paras = row_data[ci] if ci < len(row_data) else []

            if not cell_paras:
                # Empty cell — just a blank paragraph
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
            else:
                for pi, para_data in enumerate(cell_paras):
                    if pi == 0:
                        p = cell.paragraphs[0]
                    else:
                        p = cell.add_paragraph()
                    p.alignment = para_data.get("align", WD_ALIGN_PARAGRAPH.LEFT)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.line_spacing = Pt(12)
                    runs = para_data.get("runs", [])
                    for run_data in runs:
                        run = p.add_run(run_data["text"])
                        run.bold = run_data.get("bold", False) or is_header
                        run.italic = run_data.get("italic", False)
                        sz = run_data.get("size", 9)
                        run.font.size = Pt(sz if sz > 0 else 9)
                        run.font.name = "Arial"
                        if is_header:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        else:
                            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

            # Styling
            _set_cell_borders(cell)
            if is_header:
                _set_cell_shading(cell, "1F3864")
            elif ri % 2 == 0:
                _set_cell_shading(cell, "F2F5FA")

            # Cell margins
            tcPr = cell._tc.get_or_add_tcPr()
            tcMar = OxmlElement("w:tcMar")
            for side, val in [("top","60"),("bottom","60"),("left","100"),("right","100")]:
                m = OxmlElement(f"w:{side}")
                m.set(qn("w:w"), val); m.set(qn("w:type"), "dxa")
                tcMar.append(m)
            tcPr.append(tcMar)


def _build_word_doc(pdf_bytes: bytes) -> bytes:
    doc = Document()
    # Page setup: A4, narrow margins
    for sec in doc.sections:
        sec.page_width  = Inches(8.27)
        sec.page_height = Inches(11.69)
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(1.8)
        sec.right_margin  = Cm(1.8)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        page_width = pdf.pages[0].width if pdf.pages else 595.32
        margin_pts = 40  # approximate PDF margin

        for pi, page in enumerate(pdf.pages):
            if pi > 0:
                doc.add_page_break()

            blocks = _extract_page_blocks(page)

            for block in blocks:
                if block["type"] == "para":
                    # Determine font size from runs
                    sizes = [r.get("size", 10) for r in block.get("runs", []) if r.get("text","").strip()]
                    avg_size = sum(sizes) / len(sizes) if sizes else 10
                    _add_para_block(doc, block, default_size=avg_size)

                elif block["type"] == "table":
                    _add_table_block(doc, block, page_width_pts=page.width, margin_pts=margin_pts)
                    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _build_word_doc_scanned(pdf_bytes: bytes) -> bytes:
    """OCR-based Word for scanned PDFs using PyMuPDF + pytesseract."""
    import fitz, pytesseract
    from PIL import Image
    _setup_tesseract()

    doc = Document()
    for sec in doc.sections:
        sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
        sec.top_margin = sec.bottom_margin = Cm(1.8)
        sec.left_margin = sec.right_margin = Cm(1.8)

    pdf_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for pn in range(len(pdf_doc)):
        if pn > 0: doc.add_page_break()
        page = pdf_doc[pn]
        mat = fitz.Matrix(200/72, 200/72)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        ocr = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT,
                                         lang='eng', config='--oem 3 --psm 6')
        lines: dict = {}
        for i in range(len(ocr['text'])):
            txt = ocr['text'][i].strip()
            if int(ocr['conf'][i]) < 30 or not txt: continue
            key = (ocr['block_num'][i], ocr['par_num'][i], ocr['line_num'][i])
            lines.setdefault(key, []).append(txt)
        for key in sorted(lines):
            line_text = " ".join(lines[key]).strip()
            if line_text:
                p = doc.add_paragraph(line_text)
                p.paragraph_format.space_after = Pt(2)
    pdf_doc.close()
    buf = io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# PREVIEW: fast HTML, no OCR
# ─────────────────────────────────────────────────────────────────────────────

def _word_preview_html(pdf_bytes: bytes, scanned: bool) -> str:
    if scanned:
        return ('<div style="padding:32px;text-align:center;color:#555;">'
                '<p style="font-size:14px;font-weight:bold;margin-bottom:10px">📄 Scanned PDF</p>'
                '<p style="font-size:11px;color:#888;">Click <b>Convert to Word</b> to run OCR and get a fully editable .docx</p>'
                '</div>')
    parts = []
    with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
        for pn, page in enumerate(pdf.pages, 1):
            parts.append(f'<p style="color:#aaa;font-size:9px;margin:10px 0 3px;border-top:1px solid #eee;padding-top:6px">Page {pn}</p>')
            blocks = _extract_page_blocks(page)
            for block in blocks:
                if block["type"] == "para":
                    runs = block.get("runs", [])
                    if not any(r.get("text","").strip() for r in runs): continue
                    align = "center" if block.get("align") == WD_ALIGN_PARAGRAPH.CENTER else "left"
                    content = ""
                    for r in runs:
                        t = html_module.escape(r.get("text",""))
                        sz = r.get("size", 10)
                        style = f"font-size:{sz}px;"
                        if r.get("bold"): style += "font-weight:bold;"
                        if r.get("italic"): style += "font-style:italic;"
                        content += f'<span style="{style}">{t}</span>'
                    parts.append(f'<p style="margin:0 0 3px;text-align:{align}">{content}</p>')
                elif block["type"] == "table":
                    rows = block["rows"]
                    if not rows: continue
                    parts.append('<table style="border-collapse:collapse;width:100%;font-size:9px;margin:6px 0">')
                    for ri, row in enumerate(rows):
                        parts.append('<tr>')
                        for ci, cell_paras in enumerate(row):
                            cell_text = " ".join(
                                " ".join(r.get("text","") for r in p.get("runs",[]))
                                for p in cell_paras
                            ).strip()
                            is_hdr = ri == 0
                            if is_hdr:
                                parts.append(f'<th style="background:#1F3864;color:#fff;border:1px solid #aaa;padding:4px 6px;text-align:left">{html_module.escape(cell_text)}</th>')
                            else:
                                bg = "#F2F5FA" if ri%2==0 else "#fff"
                                parts.append(f'<td style="background:{bg};border:1px solid #ddd;padding:3px 6px;vertical-align:top">{html_module.escape(cell_text)}</td>')
                        parts.append('</tr>')
                    parts.append('</table>')
    return "\n".join(parts)


# ─────────────────────────────────────────────────────────────────────────────
# EXCEL — OCR-based table extraction preserving PDF layout exactly
# ─────────────────────────────────────────────────────────────────────────────

def _get_col_boundaries(drawings):
    """Detect column x-boundaries from vertical lines in PDF drawings."""
    v_lines = []
    for d in drawings:
        r = d.get("rect")
        if r and r.width < 4 and r.height > 20:
            v_lines.append(r.x0)
    if not v_lines:
        return None
    sorted_xs = sorted(set(round(x, 1) for x in v_lines))
    merged = [sorted_xs[0]]
    for x in sorted_xs[1:]:
        if x - merged[-1] > 3:
            merged.append(x)
    return merged if len(merged) >= 2 else None


def _get_row_boundaries(drawings, min_width=150, min_height=8, tol=3):
    """Detect row y-boundaries from horizontal filled rectangles."""
    rects = [
        d["rect"] for d in drawings
        if d.get("rect") and d["rect"].width > min_width and d["rect"].height > min_height
    ]
    if not rects:
        return []
    ys = set()
    for r in rects:
        ys.add(r.y0)
        ys.add(r.y1)
    ys = sorted(ys)
    merged = [ys[0]]
    for y in ys[1:]:
        if y - merged[-1] > tol:
            merged.append(y)
    return merged


def _ocr_cell_text(img, x0, y0, x1, y1, scale):
    """OCR a cell region from a page image."""
    import re as _re
    try:
        from PIL import Image as PILImage
        import pytesseract
        pad = 4
        sx0 = max(0, int(x0 * scale) - pad)
        sy0 = max(0, int(y0 * scale) - pad)
        sx1 = int(x1 * scale) + pad
        sy1 = int(y1 * scale) + pad
        crop = img.crop((sx0, sy0, sx1, sy1))
        if crop.size[0] < 8 or crop.size[1] < 8:
            return ""
        w, h = crop.size
        crop_up = crop.resize((w * 2, h * 2), PILImage.LANCZOS)
        text = pytesseract.image_to_string(crop_up, config="--psm 6 --oem 3")
        return _re.sub(r"\s+", " ", text).strip()
    except Exception:
        return ""


def _detect_cell_fill(drawings, x0, y0, x1, y1):
    """Return hex fill color if a filled rect covers this cell, else None."""
    for d in drawings:
        r = d.get("rect")
        fill = d.get("fill")
        if r and fill and isinstance(fill, (list, tuple)) and len(fill) >= 3:
            # Check overlap
            if r.x0 <= x0 + 5 and r.x1 >= x1 - 5 and r.y0 <= y0 + 2 and r.y1 >= y1 - 2:
                rgb = tuple(int(c * 255) for c in fill[:3])
                return "%02X%02X%02X" % rgb
    return None


def _assign_word_to_cell(wx, wy, col_xs, row_ys, scale):
    """Map pixel coords of a word center back to (row_idx, col_idx) in the table grid."""
    px, py = wx / scale, wy / scale
    ci = next((i for i in range(len(col_xs)-1) if col_xs[i] <= px <= col_xs[i+1]), None)
    ri = next((i for i in range(len(row_ys)-1) if row_ys[i] <= py <= row_ys[i+1]), None)
    return ri, ci


def _extract_tables_for_excel(pdf_bytes: bytes):
    """
    Extract tables from PDF using full-page OCR word-position mapping.
    Each OCR word is placed into the correct cell based on its (x,y) position
    relative to the table grid detected from vector drawing lines.
    """
    import fitz
    import pytesseract
    from PIL import Image as PILImage

    SCALE = 3.0
    result = []

    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    for pn in range(len(doc)):
        page = doc[pn]
        drawings = page.get_drawings()

        col_xs = _get_col_boundaries(drawings)
        row_ys = _get_row_boundaries(drawings)

        if not col_xs or len(col_xs) < 2 or len(row_ys) < 2:
            # Fallback: pdfplumber for text-based PDFs
            import pdfplumber
            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf_pl:
                pl_page = pdf_pl.pages[pn]
                for tn, ft in enumerate(pl_page.find_tables(), 1):
                    raw = ft.extract()
                    if not raw:
                        continue
                    nc = max(len(r) for r in raw)
                    rows = [
                        [_clean(c) for c in (r + [None] * (nc - len(r)))]
                        for r in raw if any(r)
                    ]
                    if rows:
                        result.append({
                            "label": f"Page {pn+1} — Table {tn}",
                            "rows": rows,
                            "col_widths_pts": None,
                            "cell_fills": None,
                            "header_row_idx": 0,
                        })
            continue

        n_rows = len(row_ys) - 1
        n_cols = len(col_xs) - 1

        # Render full page in-memory (no temp files — avoids Windows file-lock)
        pix = page.get_pixmap(matrix=fitz.Matrix(SCALE, SCALE))
        img = PILImage.open(io.BytesIO(pix.tobytes("png")))
        img.load()

        # Full-page OCR with word-level bounding boxes
        ocr_data = pytesseract.image_to_data(
            img, output_type=pytesseract.Output.DICT, config="--oem 3 --psm 6"
        )

        # Build empty word grid
        grid = [[[] for _ in range(n_cols)] for _ in range(n_rows)]

        for i in range(len(ocr_data["text"])):
            word = ocr_data["text"][i].strip()
            try:
                conf = int(ocr_data["conf"][i])
            except (ValueError, TypeError):
                conf = 0
            if not word or conf < 25:
                continue
            # Word center in pixels
            wx = ocr_data["left"][i] + ocr_data["width"][i] // 2
            wy = ocr_data["top"][i] + ocr_data["height"][i] // 2
            ri, ci = _assign_word_to_cell(wx, wy, col_xs, row_ys, SCALE)
            if ri is not None and ci is not None:
                grid[ri][ci].append(word)

        # Convert word lists to strings and collect fill colors
        col_widths_pts = [col_xs[i+1] - col_xs[i] for i in range(n_cols)]
        rows_out = []
        cell_fills_out = []

        for ri in range(n_rows):
            row_texts = []
            row_fills = []
            for ci in range(n_cols):
                text = " ".join(grid[ri][ci])
                x0, x1 = col_xs[ci], col_xs[ci+1]
                y0, y1 = row_ys[ri], row_ys[ri+1]
                fill = _detect_cell_fill(drawings, x0, y0, x1, y1)
                row_texts.append(text)
                row_fills.append(fill)
            rows_out.append(row_texts)
            cell_fills_out.append(row_fills)

        if not rows_out:
            continue

        result.append({
            "label": f"Page {pn+1} — Table",
            "rows": rows_out,
            "col_widths_pts": col_widths_pts,
            "cell_fills": cell_fills_out,
            "header_row_idx": 0,
        })

    doc.close()
    return result



def _pdf_to_excel(pdf_bytes: bytes) -> bytes:
    """Build an Excel workbook that mirrors the PDF table layout exactly."""
    thin = Side(style="thin", color="AAAAAA")
    medium = Side(style="medium", color="4F4F4F")
    border_thin = Border(left=thin, right=thin, top=thin, bottom=thin)
    border_header = Border(left=medium, right=medium, top=medium, bottom=medium)

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "PDF Tables"

    cur_row = 1

    for tbl in _extract_tables_for_excel(pdf_bytes):
        rows = tbl["rows"]
        col_widths_pts = tbl["col_widths_pts"]
        cell_fills = tbl["cell_fills"]
        header_row_idx = tbl.get("header_row_idx", 0)

        if not rows:
            continue

        n_cols = len(rows[0])

        # Table label banner
        label_cell = ws.cell(row=cur_row, column=1, value=tbl["label"])
        label_cell.font = Font(name="Arial", bold=True, size=9, color="1F3864")
        label_cell.fill = PatternFill("solid", fgColor="D9E1F2")
        label_cell.alignment = Alignment(horizontal="left", vertical="center", indent=1)
        ws.row_dimensions[cur_row].height = 16
        # Merge label across all columns
        if n_cols > 1:
            ws.merge_cells(
                start_row=cur_row, start_column=1,
                end_row=cur_row, end_column=n_cols
            )
        cur_row += 1

        # Compute column widths for Excel (convert PDF points to characters ~7px each)
        if col_widths_pts:
            pts_to_chars = lambda pts: max(8, min(60, int(pts * 0.18)))
            excel_col_widths = [pts_to_chars(w) for w in col_widths_pts]
        else:
            # Auto-width from content
            excel_col_widths = None

        # Write rows
        for ri, row in enumerate(rows):
            is_header = (ri == header_row_idx)
            excel_row = cur_row + ri
            ws.row_dimensions[excel_row].height = None  # auto

            for ci, cell_text in enumerate(row):
                cell = ws.cell(row=excel_row, column=ci + 1, value=cell_text or "")

                # Determine fill color
                fill_hex = None
                if cell_fills and ri < len(cell_fills) and ci < len(cell_fills[ri]):
                    fill_hex = cell_fills[ri][ci]

                # Header row: dark blue background, white bold text
                if is_header:
                    cell.font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
                    cell.fill = PatternFill("solid", fgColor="1F3864")
                    cell.border = border_header
                    cell.alignment = Alignment(
                        horizontal="left", vertical="center",
                        wrap_text=True
                    )
                    # Set header row height
                    ws.row_dimensions[excel_row].height = 28
                else:
                    cell.font = Font(name="Arial", size=9, color="1A1A1A")
                    cell.border = border_thin
                    cell.alignment = Alignment(
                        wrap_text=True, vertical="top", horizontal="left"
                    )
                    # Apply PDF fill color if detected, else alternate row shading
                    if fill_hex and fill_hex not in ("BFBFBF", "C0C0C0", "BEBEBE"):
                        cell.fill = PatternFill("solid", fgColor=fill_hex)
                    elif (ri % 2) == 1:
                        cell.fill = PatternFill("solid", fgColor="EEF2F7")
                    else:
                        cell.fill = PatternFill("solid", fgColor="FFFFFF")

                    # Estimate row height by content length
                    if cell_text:
                        lines = max(1, len(cell_text) // 60 + cell_text.count("\n") + 1)
                        cur_h = ws.row_dimensions[excel_row].height or 15
                        ws.row_dimensions[excel_row].height = max(cur_h, min(120, lines * 14))

        cur_row += len(rows)

        # Set column widths
        if excel_col_widths:
            for ci, w in enumerate(excel_col_widths):
                ws.column_dimensions[get_column_letter(ci + 1)].width = w
        else:
            # Auto-size from content
            cw = {}
            for r in ws.iter_rows(min_row=cur_row - len(rows), max_row=cur_row - 1):
                for cell in r:
                    if cell.value:
                        first_line = str(cell.value).split("\n")[0]
                        cw[cell.column] = max(cw.get(cell.column, 8), min(55, len(first_line) + 3))
            for col_idx, w in cw.items():
                ws.column_dimensions[get_column_letter(col_idx)].width = w

        # Spacer row
        ws.row_dimensions[cur_row].height = 8
        cur_row += 1

    # Freeze top rows up to first data header
    ws.freeze_panes = "A2"
    ws.page_setup.orientation = "landscape"
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.read()


# ─────────────────────────────────────────────────────────────────────────────
# ROUTES
# ─────────────────────────────────────────────────────────────────────────────

def _get_pdf():
    if "file" not in request.files: return None, (jsonify({"error":"No file"}), 400)
    f = request.files["file"]; b = f.read()
    if not b: return None, (jsonify({"error":"Empty"}), 400)
    return b, f.filename

@app.route("/api/health")
def health(): return jsonify({"status":"ok"})

@app.route("/api/preview/word", methods=["POST"])
def route_preview_word():
    b, info = _get_pdf()
    if b is None: return info
    try:
        scanned = _is_scanned(b)
        return jsonify({"html": _word_preview_html(b, scanned), "scanned": scanned})
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/api/convert/word", methods=["POST"])
def route_convert_word():
    b, fname = _get_pdf()
    if b is None: return fname
    try:
        docx = _build_word_doc_scanned(b) if _is_scanned(b) else _build_word_doc(b)
        name = str(fname).rsplit(".",1)[0] + ".docx"
        return send_file(io.BytesIO(docx),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True, download_name=name)
    except Exception as e:
        import traceback; traceback.print_exc()
        return jsonify({"error": str(e)}), 500

@app.route("/api/preview/excel", methods=["POST"])
def route_preview_excel():
    b, _ = _get_pdf()
    if b is None: return _
    try:
        tables = _extract_tables_for_excel(b)
        out = []
        for t in tables:
            rows = t.get("rows", [])
            fills = t.get("cell_fills", None)
            col_widths = t.get("col_widths_pts", None)
            header_idx = t.get("header_row_idx", 0)
            out.append({
                "label": t["label"],
                "headers": rows[header_idx] if rows else [],
                "rows": rows[header_idx+1:] if rows else [],
                "all_rows": rows,
                "cell_fills": fills,
                "col_widths_pts": col_widths,
                "header_row_idx": header_idx,
            })
        return jsonify({"tables": out})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route("/api/convert/excel", methods=["POST"])
def route_convert_excel():
    b, fname = _get_pdf()
    if b is None: return fname
    try:
        xlsx = _pdf_to_excel(b)
        name = str(fname).rsplit(".",1)[0] + ".xlsx"
        return send_file(io.BytesIO(xlsx),
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True, download_name=name)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == "__main__":
    port = int(os.environ.get("BACKEND_PORT", 5050))
    print(f"✅  PDF backend running on http://localhost:{port}")
    app.run(host="0.0.0.0", port=port, debug=False)
